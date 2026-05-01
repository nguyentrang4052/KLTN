# import fitz  # PyMuPDF
# from docx import Document
# from typing import List, Dict, Any
# import re
# import structlog
# from io import BytesIO

# logger = structlog.get_logger()

# class CVProcessor:
#     def __init__(self):
#         self.chunk_size = 1500
#         self.chunk_overlap = 300
#         self.max_text_length = 50000  # KHÔNG ĐƯỢC THIẾU
    
#     async def extract_text(self, file_bytes: bytes, mime_type: str) -> str:
#         """Extract text với giới hạn kích thước"""
#         try:
#             # Kiểm tra file rỗng
#             if not file_bytes or len(file_bytes) == 0:
#                 logger.error("Empty file provided")
#                 raise ValueError("File trống")
            
#             # Kiểm tra kích thước file
#             if len(file_bytes) > 10 * 1024 * 1024:
#                 raise ValueError("File quá lớn (>10MB)")
            
#             logger.info(f"Extracting text from file type: {mime_type}, size: {len(file_bytes)} bytes")
            
#             # Xử lý theo loại file
#             if mime_type == "application/pdf":
#                 text = await self._extract_pdf(file_bytes)
#             elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
#                 text = await self._extract_docx(file_bytes)
#             else:
#                 raise ValueError(f"Không hỗ trợ định dạng file: {mime_type}")
            
#             # Kiểm tra text rỗng
#             if not text or len(text.strip()) < 50:
#                 logger.warning(f"Extracted text too short: {len(text)} chars")
#                 raise ValueError("Không đọc được nội dung từ file")
            
#             # Giới hạn độ dài
#             if len(text) > self.max_text_length:
#                 text = text[:self.max_text_length] + "\n...[Đã cắt ngắn]..."
            
#             logger.info(f"Successfully extracted {len(text)} characters")
#             return text
            
#         except Exception as e:
#             logger.error(f"extract_text_error: {str(e)}")
#             raise
    
#     async def _extract_pdf(self, file_bytes: bytes) -> str:
#         """Extract text từ PDF"""
#         text = ""
#         try:
#             # Mở PDF từ bytes
#             doc = fitz.open(stream=file_bytes, filetype="pdf")
            
#             for i, page in enumerate(doc):
#                 if i >= 50:  # Giới hạn 50 trang
#                     break
#                 page_text = page.get_text()
#                 if page_text:
#                     text += page_text + "\n"
            
#             doc.close()
            
#             if not text.strip():
#                 logger.warning("PDF has no extractable text (might be scanned image)")
#                 raise ValueError("PDF không có text (có thể là file scan)")
            
#             return self._clean_text(text)
            
#         except Exception as e:
#             logger.error(f"PDF extraction error: {str(e)}")
#             raise ValueError(f"Không đọc được file PDF: {str(e)}")
    
    
#     async def _extract_docx(self, file_bytes: bytes) -> str:
#         """Extract text từ DOCX"""
#         try:
#             doc = Document(BytesIO(file_bytes))
#             text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
#             # Cũng lấy text từ bảng nếu có
#             for table in doc.tables:
#                 for row in table.rows:
#                     for cell in row.cells:
#                         if cell.text.strip():
#                             text += "\n" + cell.text
            
#             if not text.strip():
#                 logger.warning("DOCX has no extractable text")
#                 raise ValueError("DOCX không có nội dung text")
            
#             return self._clean_text(text)
            
#         except Exception as e:
#             logger.error(f"DOCX extraction error: {str(e)}")
#             raise ValueError(f"Không đọc được file DOCX: {str(e)}")
    
#     def create_chunks(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
#         sentences = re.split(r'(?<=[.!?])\s+', text)
#         chunks = []
#         current_chunk = []
#         current_length = 0
        
#         for sentence in sentences:
#             sentence_length = len(sentence)
#             if current_length + sentence_length > self.chunk_size and current_chunk:
#                 chunks.append({
#                     "content": " ".join(current_chunk),
#                     "metadata": {**metadata, "chunk_index": len(chunks)}
#                 })
#                 overlap = current_chunk[-3:] if len(current_chunk) > 3 else []
#                 current_chunk = overlap + [sentence]
#                 current_length = sum(len(s) for s in overlap) + sentence_length
#             else:
#                 current_chunk.append(sentence)
#                 current_length += sentence_length
        
#         if current_chunk:
#             chunks.append({
#                 "content": " ".join(current_chunk),
#                 "metadata": {**metadata, "chunk_index": len(chunks)}
#             })
        
#         return chunks
    
#     def extract_sections(self, text: str) -> Dict[str, str]:
#         sections = {}
#         patterns = {
#             "experience": r'((kinh nghiệm|experience)[\s\S]*?)(?=(học vấn|education|$))',
#             "education": r'((học vấn|education)[\s\S]*?)(?=(kinh nghiệm|experience|$))',
#             "skills": r'((kỹ năng|skills)[\s\S]*?)(?=(kinh nghiệm|experience|$))',
#         }
#         for section, pattern in patterns.items():
#             match = re.search(pattern, text, re.IGNORECASE)
#             if match:
#                 sections[section] = match.group(1).strip()[:2000]
#         return sections
    
#     def _clean_text(self, text: str) -> str:
#         """Làm sạch text"""
#         if not text:
#             return ""
#         # Loại bỏ khoảng trắng thừa
#         text = re.sub(r'\s+', ' ', text)
#         # Loại bỏ các ký tự đặc biệt không cần thiết
#         text = re.sub(r'[^\w\s\.,;:?!@#$%^&*()\-+=/\\|<>\[\]{}~`\'"áàảãạăắằẳẵặâấầẩẫậđéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵ]', ' ', text)
#         text = re.sub(r'\s+', ' ', text).strip()
#         return text


import fitz  # PyMuPDF
from docx import Document
from typing import List, Dict, Any
import re
import structlog
from io import BytesIO
import asyncio
from concurrent.futures import ThreadPoolExecutor
import base64

# OCR imports
try:
    import pytesseract
    from PIL import Image
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
    print("Warning: pytesseract not installed")

logger = structlog.get_logger()

class CVProcessor:
    def __init__(self):
        self.chunk_size = 1500
        self.chunk_overlap = 300
        self.max_text_length = 50000
        self.executor = ThreadPoolExecutor(max_workers=2)
        
        # Cấu hình Tesseract path cho Windows
        if PYTESSERACT_AVAILABLE:
            import platform
            if platform.system() == "Windows":
                # Đường dẫn mặc định của Tesseract trên Windows
                tesseract_paths = [
                    r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                    r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                ]
                for path in tesseract_paths:
                    import os
                    if os.path.exists(path):
                        pytesseract.pytesseract.tesseract_cmd = path
                        logger.info(f"Tesseract found at: {path}")
                        break
            # Ngôn ngữ: tiếng Việt + English
            self.tesseract_lang = "vie+eng"
            logger.info("OCR initialized with Tesseract")
    
    async def extract_text(self, file_bytes: bytes, mime_type: str) -> str:
        """Extract text với hỗ trợ OCR cho PDF scan"""
        try:
            if not file_bytes or len(file_bytes) == 0:
                raise ValueError("File trống")
            
            if len(file_bytes) > 10 * 1024 * 1024:
                raise ValueError("File quá lớn (>10MB)")
            
            logger.info(f"Extracting text from: {mime_type}, size: {len(file_bytes)} bytes")
            
            if mime_type == "application/pdf":
                text = await self._extract_pdf_with_ocr(file_bytes)
            elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                text = await self._extract_docx(file_bytes)
            else:
                raise ValueError(f"Không hỗ trợ định dạng: {mime_type}")
            
            if not text or len(text.strip()) < 50:
                raise ValueError("Không thể trích xuất nội dung từ file")
            
            # Giới hạn độ dài
            if len(text) > self.max_text_length:
                text = text[:self.max_text_length] + "\n...[Đã cắt ngắn]..."
            
            logger.info(f"Extracted {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"extract_text_error: {str(e)}")
            raise
    
    async def _extract_pdf_with_ocr(self, file_bytes: bytes) -> str:
        """Extract text từ PDF, dùng OCR nếu không có text layer"""
        text = ""
        
        try:
            # Mở PDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            
            # Thử lấy text trực tiếp
            for page_num, page in enumerate(doc):
                if page_num >= 50:
                    break
                page_text = page.get_text()
                if page_text and len(page_text.strip()) > 50:
                    text += page_text + "\n"
            
            doc.close()
            
            # Nếu text quá ít (file scan), chuyển sang OCR
            if len(text.strip()) < 100:
                logger.info("Text layer too short, using OCR for scanned PDF...")
                text = await self._ocr_pdf(file_bytes)
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise ValueError(f"Không đọc được file PDF: {str(e)}")
    
    async def _ocr_pdf(self, file_bytes: bytes) -> str:
        """Sử dụng OCR để đọc text từ PDF scan"""
        if not PYTESSERACT_AVAILABLE:
            raise ValueError("OCR không khả dụng. Vui lòng cài đặt pytesseract và Tesseract OCR.")
        
        def ocr_pdf_sync(file_bytes):
            text_parts = []
            try:
                # Mở PDF
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                
                for page_num in range(min(len(doc), 20)):  # Giới hạn 20 trang
                    page = doc[page_num]
                    
                    # Tăng độ phân giải để OCR tốt hơn
                    zoom = 2.5
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Chuyển sang ảnh
                    img_data = pix.tobytes("png")
                    image = Image.open(BytesIO(img_data))
                    
                    # OCR với cấu hình tối ưu cho tiếng Việt
                    custom_config = r'--oem 3 --psm 6 -l vie+eng'
                    page_text = pytesseract.image_to_string(image, config=custom_config)
                    
                    if page_text and len(page_text.strip()) > 10:
                        text_parts.append(page_text)
                        logger.info(f"OCR page {page_num + 1}: {len(page_text)} chars")
                    else:
                        logger.warning(f"OCR page {page_num + 1} returned empty")
                
                doc.close()
                return "\n".join(text_parts)
                
            except Exception as e:
                logger.error(f"OCR error: {str(e)}")
                raise ValueError(f"OCR thất bại: {str(e)}")
        
        # Chạy OCR trong thread pool
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(self.executor, ocr_pdf_sync, file_bytes)
        
        if not text or len(text.strip()) < 100:
            raise ValueError("OCR không thể đọc được nội dung từ file scan")
        
        return text
    
    async def _extract_docx(self, file_bytes: bytes) -> str:
        """Extract text từ DOCX"""
        try:
            doc = Document(BytesIO(file_bytes))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            # Lấy text từ bảng
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += "\n" + cell.text
            
            if not text.strip():
                raise ValueError("DOCX không có nội dung text")
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise ValueError(f"Không đọc được file DOCX: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Làm sạch text"""
        if not text:
            return ""
        # Loại bỏ khoảng trắng thừa
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        return text
    

    def create_chunks(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            if current_length + sentence_length > self.chunk_size and current_chunk:
                chunks.append({
                    "content": " ".join(current_chunk),
                    "metadata": {**metadata, "chunk_index": len(chunks)}
                })
                overlap = current_chunk[-3:] if len(current_chunk) > 3 else []
                current_chunk = overlap + [sentence]
                current_length = sum(len(s) for s in overlap) + sentence_length
            else:
                current_chunk.append(sentence)
                current_length += sentence_length
        
        if current_chunk:
            chunks.append({
                "content": " ".join(current_chunk),
                "metadata": {**metadata, "chunk_index": len(chunks)}
            })
        
        return chunks
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        sections = {}
        patterns = {
            "experience": r'((kinh nghiệm|experience)[\s\S]*?)(?=(học vấn|education|$))',
            "education": r'((học vấn|education)[\s\S]*?)(?=(kinh nghiệm|experience|$))',
            "skills": r'((kỹ năng|skills)[\s\S]*?)(?=(kinh nghiệm|experience|$))',
        }
        for section, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                sections[section] = match.group(1).strip()[:2000]
        return sections